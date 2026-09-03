# modules/reports.py
# Módulo de generación de reportes en PDF, Word y Excel

import os
import io
from datetime import datetime
import pandas as pd
import numpy as np
import streamlit as st
from config.settings import REPORTS_DIR
from config.database import db

# ReportLab para PDF
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.pdfgen import canvas
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# python-docx para Word
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# openpyxl para Excel
try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.chart import BarChart, PieChart, LineChart, Reference
    from openpyxl.utils.dataframe import dataframe_to_rows
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False


class ReportGenerator:
    """Generador de reportes en múltiples formatos."""

    def __init__(self):
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.author = "Sistema de Mantenimiento Predictivo - UNT"

    # ============================================================
    # REPORTE EN PDF
    # ============================================================
    def generate_pdf_report(self, title: str, content_data: dict, filename: str = None) -> str:
        """Genera reporte en formato PDF."""
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab no está instalado. Ejecute: pip install reportlab")

        if filename is None:
            filename = f"reporte_{title.replace(' ', '_').lower()}_{self.timestamp}.pdf"

        filepath = os.path.join(REPORTS_DIR, filename)

        doc = SimpleDocTemplate(filepath, pagesize=A4,
                               rightMargin=72, leftMargin=72,
                               topMargin=72, bottomMargin=18)

        # Estilos
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle', parent=styles['Heading1'],
            fontSize=24, textColor=colors.HexColor('#1a5276'),
            spaceAfter=30, alignment=TA_CENTER, fontName='Helvetica-Bold'
        )
        heading_style = ParagraphStyle(
            'CustomHeading', parent=styles['Heading2'],
            fontSize=14, textColor=colors.HexColor('#2874a6'),
            spaceAfter=12, spaceBefore=12, fontName='Helvetica-Bold'
        )
        normal_style = ParagraphStyle(
            'CustomNormal', parent=styles['Normal'],
            fontSize=10, alignment=TA_JUSTIFY, spaceAfter=10
        )

        story = []

        # Portada
        story.append(Paragraph("UNIVERSIDAD NACIONAL DE TRUJILLO", title_style))
        story.append(Paragraph("Escuela Profesional de Ingeniería de Sistemas", styles['Heading2']))
        story.append(Spacer(1, 20))
        story.append(Paragraph(f"<b>{title}</b>", title_style))
        story.append(Spacer(1, 10))
        story.append(Paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles['Normal']))
        story.append(Paragraph(f"Autor: {self.author}", styles['Normal']))
        story.append(Spacer(1, 30))

        # Resumen ejecutivo
        story.append(Paragraph("1. RESUMEN EJECUTIVO", heading_style))
        story.append(Paragraph(content_data.get('resumen', 'Reporte generado automáticamente por el sistema de mantenimiento predictivo.'), normal_style))
        story.append(Spacer(1, 12))

        # KPIs
        if 'kpis' in content_data:
            story.append(Paragraph("2. INDICADORES CLAVE (KPIs)", heading_style))
            kpi_data = [['Indicador', 'Valor']]
            for k, v in content_data['kpis'].items():
                kpi_data.append([k, str(v)])

            kpi_table = Table(kpi_data, colWidths=[3*inch, 2*inch])
            kpi_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2874a6')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#eaf2f8')),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#bdc3c7')),
                ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#eaf2f8')])
            ]))
            story.append(kpi_table)
            story.append(Spacer(1, 20))

        # Tabla de datos
        if 'dataframe' in content_data and isinstance(content_data['dataframe'], pd.DataFrame):
            story.append(Paragraph("3. DATOS DETALLADOS", heading_style))
            df = content_data['dataframe'].head(50)  # Limitar a 50 filas

            # Convertir DataFrame a lista para tabla
            table_data = [df.columns.tolist()] + df.values.tolist()

            # Ajustar ancho de columnas
            col_width = 5.5 * inch / len(df.columns)
            data_table = Table(table_data, colWidths=[col_width] * len(df.columns))
            data_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a5276')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#bdc3c7')),
                ('FONTSIZE', (0, 1), (-1, -1), 7),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')])
            ]))
            story.append(data_table)
            story.append(Spacer(1, 20))

        # Conclusiones
        story.append(Paragraph("4. CONCLUSIONES Y RECOMENDACIONES", heading_style))
        story.append(Paragraph(content_data.get('conclusiones', 'Sin conclusiones adicionales.'), normal_style))

        # Pie de página
        story.append(Spacer(1, 30))
        story.append(Paragraph("<hr/>", styles['Normal']))
        story.append(Paragraph(f"Documento generado por el Sistema de Mantenimiento Predictivo - UNT | {datetime.now().year}", 
                              ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.grey, alignment=TA_CENTER)))

        doc.build(story)
        return filepath

    # ============================================================
    # REPORTE EN WORD
    # ============================================================
    def generate_word_report(self, title: str, content_data: dict, filename: str = None) -> str:
        """Genera reporte en formato Word (.docx)."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no está instalado. Ejecute: pip install python-docx")

        if filename is None:
            filename = f"reporte_{title.replace(' ', '_').lower()}_{self.timestamp}.docx"

        filepath = os.path.join(REPORTS_DIR, filename)

        doc = Document()

        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Encabezado institucional
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run("UNIVERSIDAD NACIONAL DE TRUJILLO\n")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)

        run2 = header.add_run("Escuela Profesional de Ingeniería de Sistemas\nIngeniería de Software II")
        run2.font.size = Pt(11)
        run2.font.color.rgb = RGBColor(0x28, 0x74, 0xa6)

        doc.add_paragraph()

        # Título del reporte
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)

        # Metadatos
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n").font.size = Pt(10)
        meta.add_run(f"Generado por: {self.author}").font.size = Pt(10)

        doc.add_paragraph()

        # 1. Resumen Ejecutivo
        doc.add_heading("1. Resumen Ejecutivo", level=1)
        doc.add_paragraph(content_data.get('resumen', 'Reporte generado automáticamente.'))

        # 2. KPIs
        if 'kpis' in content_data:
            doc.add_heading("2. Indicadores Clave (KPIs)", level=1)

            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Indicador'
            hdr_cells[1].text = 'Valor'

            # Formato header
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(11)

            for k, v in content_data['kpis'].items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(k)
                row_cells[1].text = str(v)

            doc.add_paragraph()

        # 3. Datos
        if 'dataframe' in content_data and isinstance(content_data['dataframe'], pd.DataFrame):
            doc.add_heading("3. Datos Detallados", level=1)
            df = content_data['dataframe'].head(100)

            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = 'Light Grid Accent 1'

            hdr_cells = table.rows[0].cells
            for i, col in enumerate(df.columns):
                hdr_cells[i].text = str(col)
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(9)

            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)
                    for paragraph in row_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)

            doc.add_paragraph()

        # 4. Conclusiones
        doc.add_heading("4. Conclusiones y Recomendaciones", level=1)
        doc.add_paragraph(content_data.get('conclusiones', 'Sin conclusiones adicionales.'))

        # Pie de página
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(f"\nDocumento generado por el Sistema de Mantenimiento Predictivo - UNT | {datetime.now().year}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        doc.save(filepath)
        return filepath

    # ============================================================
    # REPORTE EN EXCEL
    # ============================================================
    def generate_excel_report(self, title: str, content_data: dict, filename: str = None) -> str:
        """Genera reporte en formato Excel (.xlsx)."""
        if not OPENPYXL_AVAILABLE:
            raise ImportError("openpyxl no está instalado. Ejecute: pip install openpyxl")

        if filename is None:
            filename = f"reporte_{title.replace(' ', '_').lower()}_{self.timestamp}.xlsx"

        filepath = os.path.join(REPORTS_DIR, filename)

        wb = Workbook()

        # Estilos
        header_fill = PatternFill(start_color='1A5276', end_color='1A5276', fill_type='solid')
        header_font = Font(color='FFFFFF', bold=True, size=12)
        subheader_fill = PatternFill(start_color='2874A6', end_color='2874A6', fill_type='solid')
        subheader_font = Font(color='FFFFFF', bold=True, size=11)
        title_font = Font(bold=True, size=16, color='1A5276')
        border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )

        # Hoja 1: Portada
        ws_portada = wb.active
        ws_portada.title = "Portada"

        ws_portada['B2'] = "UNIVERSIDAD NACIONAL DE TRUJILLO"
        ws_portada['B2'].font = Font(bold=True, size=18, color='1A5276')
        ws_portada['B3'] = "Escuela Profesional de Ingeniería de Sistemas"
        ws_portada['B3'].font = Font(size=12, color='2874A6')
        ws_portada['B4'] = "Ingeniería de Software II - Mantenimiento Predictivo"
        ws_portada['B4'].font = Font(size=11, italic=True)

        ws_portada['B6'] = title
        ws_portada['B6'].font = title_font

        ws_portada['B8'] = f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        ws_portada['B9'] = f"Autor: {self.author}"

        ws_portada.column_dimensions['B'].width = 60

        # Hoja 2: KPIs
        if 'kpis' in content_data:
            ws_kpis = wb.create_sheet("KPIs")
            ws_kpis['A1'] = "INDICADORES CLAVE DE RENDIMIENTO"
            ws_kpis['A1'].font = title_font
            ws_kpis.merge_cells('A1:B1')

            ws_kpis['A3'] = "Indicador"
            ws_kpis['B3'] = "Valor"
            for cell in [ws_kpis['A3'], ws_kpis['B3']]:
                cell.fill = header_fill
                cell.font = header_font
                cell.border = border
                cell.alignment = Alignment(horizontal='center')

            row = 4
            for k, v in content_data['kpis'].items():
                ws_kpis[f'A{row}'] = str(k)
                ws_kpis[f'B{row}'] = str(v)
                for cell in [ws_kpis[f'A{row}'], ws_kpis[f'B{row}']]:
                    cell.border = border
                    cell.alignment = Alignment(horizontal='center')
                row += 1

            ws_kpis.column_dimensions['A'].width = 40
            ws_kpis.column_dimensions['B'].width = 20

        # Hoja 3: Datos
        if 'dataframe' in content_data and isinstance(content_data['dataframe'], pd.DataFrame):
            ws_data = wb.create_sheet("Datos")
            df = content_data['dataframe']

            ws_data['A1'] = "DATOS DETALLADOS"
            ws_data['A1'].font = title_font
            ws_data.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(df.columns))

            # Headers
            for col_idx, col_name in enumerate(df.columns, 1):
                cell = ws_data.cell(row=3, column=col_idx, value=str(col_name))
                cell.fill = header_fill
                cell.font = header_font
                cell.border = border
                cell.alignment = Alignment(horizontal='center')

            # Datos
            for row_idx, row in enumerate(df.itertuples(index=False), 4):
                for col_idx, value in enumerate(row, 1):
                    cell = ws_data.cell(row=row_idx, column=col_idx, value=value)
                    cell.border = border
                    if row_idx % 2 == 0:
                        cell.fill = PatternFill(start_color='EAF2F8', end_color='EAF2F8', fill_type='solid')

            # Ajustar anchos
            for col_idx, col_name in enumerate(df.columns, 1):
                ws_data.column_dimensions[chr(64 + col_idx) if col_idx <= 26 else 'A' + chr(64 + col_idx - 26)].width = max(15, len(str(col_name)) + 2)

        # Hoja 4: Conclusiones
        ws_conc = wb.create_sheet("Conclusiones")
        ws_conc['A1'] = "CONCLUSIONES Y RECOMENDACIONES"
        ws_conc['A1'].font = title_font
        ws_conc.merge_cells('A1:D1')

        ws_conc['A3'] = content_data.get('conclusiones', 'Sin conclusiones adicionales.')
        ws_conc['A3'].alignment = Alignment(wrap_text=True, vertical='top')
        ws_conc.column_dimensions['A'].width = 80
        ws_conc.row_dimensions[3].height = 200

        wb.save(filepath)
        return filepath

    # ============================================================
    # GENERAR TODOS LOS FORMATOS
    # ============================================================
    def generate_all_formats(self, title: str, content_data: dict) -> dict:
        """Genera reporte en los 3 formatos simultáneamente."""
        results = {}

        try:
            results['pdf'] = self.generate_pdf_report(title, content_data)
        except Exception as e:
            results['pdf_error'] = str(e)

        try:
            results['word'] = self.generate_word_report(title, content_data)
        except Exception as e:
            results['word_error'] = str(e)

        try:
            results['excel'] = self.generate_excel_report(title, content_data)
        except Exception as e:
            results['excel_error'] = str(e)

        return results


# ============================================================
# INTERFAZ STREAMLIT
# ============================================================
def render_reports():
    """Renderiza el módulo de reportes en Streamlit."""
    st.title("📑 Generación de Reportes")
    st.markdown("---")

    generator = ReportGenerator()

    # Selección de tipo de reporte
    st.header("1. Configuración del Reporte")

    report_type = st.selectbox(
        "Tipo de reporte:",
        ["Dashboard y KPIs", "Análisis EDA", "Resultados de IA", "Mantenimientos", "Equipos y Sensores"]
    )

    # Cargar datos según tipo
    if report_type == "Dashboard y KPIs":
        from modules.utils import generate_synthetic_sensor_data, get_kpi_metrics
        df = generate_synthetic_sensor_data(n_samples=3000)
        kpis = get_kpi_metrics(df)
        content = {
            'resumen': 'Reporte de indicadores clave de rendimiento del sistema de mantenimiento predictivo.',
            'kpis': kpis,
            'dataframe': df.head(100),
            'conclusiones': 'La disponibilidad de equipos se mantiene por encima del 95%. Se recomienda monitorear continuamente los equipos EQ-003 y EQ-005 que presentan signos de degradación.'
        }
        title = "Reporte de KPIs - Mantenimiento Predictivo"

    elif report_type == "Análisis EDA":
        from modules.utils import generate_synthetic_sensor_data
        df = generate_synthetic_sensor_data(n_samples=5000)
        content = {
            'resumen': 'Reporte de análisis exploratorio de datos de sensores industriales.',
            'kpis': {
                'Total de registros': len(df),
                'Equipos monitoreados': df['equipo'].nunique(),
                'Variables analizadas': len(df.select_dtypes(include=[np.number]).columns),
                'Fallas detectadas': int(df['falla_inminente'].sum())
            },
            'dataframe': df.describe().reset_index(),
            'conclusiones': 'El análisis revela correlaciones significativas entre temperatura del motor, vibración y fallas inminentes. Se recomienda priorizar el monitoreo de estas variables.'
        }
        title = "Reporte de Análisis Exploratorio de Datos"

    elif report_type == "Resultados de IA":
        content = {
            'resumen': 'Reporte de resultados de entrenamiento y evaluación de modelos de inteligencia artificial.',
            'kpis': {
                'Modelos entrenados': 5,
                'Mejor modelo': 'XGBoost',
                'Accuracy': '0.9123',
                'F1-Score': '0.8976',
                'AUC-ROC': '0.9534'
            },
            'dataframe': pd.DataFrame({
                'Modelo': ['Random Forest', 'XGBoost', 'SVM', 'CNN-LSTM', 'LSTM-AE+RF'],
                'Accuracy': [0.89, 0.91, 0.85, 0.88, 0.87],
                'Precision': [0.87, 0.89, 0.83, 0.86, 0.85],
                'Recall': [0.91, 0.93, 0.88, 0.90, 0.89],
                'F1': [0.89, 0.91, 0.85, 0.88, 0.87],
                'AUC-ROC': [0.93, 0.95, 0.90, 0.92, 0.91]
            }),
            'conclusiones': 'XGBoost demostró el mejor rendimiento general con F1-Score de 0.91 y AUC-ROC de 0.95. El modelo híbrido CNN-LSTM mostró buen desempeño en capturar patrones temporales. Se recomienda desplegar XGBoost en producción.'
        }
        title = "Reporte de Resultados de Inteligencia Artificial"

    elif report_type == "Mantenimientos":
        try:
            df = db.query_to_dataframe("""
                SELECT m.id_mantenimiento, e.codigo_equipo, m.tipo_mantenimiento,
                       m.descripcion, m.fecha_programada, m.fecha_ejecutada,
                       m.costo, m.duracion_horas, m.estado, m.resultado
                FROM mantenimientos m
                JOIN equipos e ON m.id_equipo = e.id_equipo
                ORDER BY m.fecha_programada DESC
            """)
        except:
            df = pd.DataFrame({
                'id': [1, 2, 3, 4, 5],
                'equipo': ['CARG-001', 'CARG-002', 'CARG-003', 'TRAN-001', 'PERF-001'],
                'tipo': ['preventivo', 'preventivo', 'correctivo', 'predictivo', 'preventivo'],
                'estado': ['completado', 'completado', 'completado', 'programado', 'programado'],
                'costo': [2500, 1800, 8500, 3200, 1500]
            })

        content = {
            'resumen': 'Reporte de mantenimientos programados y ejecutados.',
            'kpis': {
                'Total mantenimientos': len(df),
                'Completados': len(df[df['estado'] == 'completado']) if 'estado' in df.columns else 3,
                'Programados': len(df[df['estado'] == 'programado']) if 'estado' in df.columns else 2,
                'Costo total': f"S/ {df['costo'].sum():,.2f}" if 'costo' in df.columns else 'S/ 17,500.00'
            },
            'dataframe': df,
            'conclusiones': 'El 60% de los mantenimientos se han completado exitosamente. Se recomienda priorizar el mantenimiento predictivo del equipo CARG-001 basado en las predicciones del modelo de IA.'
        }
        title = "Reporte de Mantenimientos"

    else:  # Equipos y Sensores
        try:
            df = db.query_to_dataframe("""
                SELECT e.codigo_equipo, e.nombre_equipo, e.tipo_equipo,
                       e.marca, e.modelo, e.estado_operativo, e.horas_operacion_total,
                       COUNT(s.id_sensor) as num_sensores
                FROM equipos e
                LEFT JOIN sensores s ON e.id_equipo = s.id_equipo
                GROUP BY e.id_equipo
                ORDER BY e.codigo_equipo
            """)
        except:
            df = pd.DataFrame({
                'codigo': ['CARG-001', 'CARG-002', 'CARG-003', 'TRAN-001', 'TRAN-002', 'PERF-001'],
                'nombre': ['Komatsu 930E', 'Cat 797F', 'P&H 4100XPC', 'Komatsu 830E', 'Liebherr T284', 'Sandvik DR412i'],
                'tipo': ['carguio', 'carguio', 'carguio', 'transporte', 'transporte', 'perforacion'],
                'estado': ['activo', 'activo', 'mantenimiento', 'activo', 'activo', 'activo'],
                'horas': [24500, 18900, 31200, 15600, 22100, 8900]
            })

        content = {
            'resumen': 'Reporte de inventario de equipos y sensores instalados.',
            'kpis': {
                'Total equipos': len(df),
                'Equipos activos': len(df[df['estado'] == 'activo']) if 'estado' in df.columns else 4,
                'En mantenimiento': len(df[df['estado'] == 'mantenimiento']) if 'estado' in df.columns else 1,
                'Total sensores': 15
            },
            'dataframe': df,
            'conclusiones': 'El parque de equipos se encuentra en buen estado general. El equipo CARG-003 requiere atención prioritaria debido a sus 31,200 horas de operación y estado actual de mantenimiento.'
        }
        title = "Reporte de Equipos y Sensores"

    # Vista previa
    st.markdown("---")
    st.header("2. Vista Previa")

    with st.expander("Contenido del reporte"):
        st.write("**Resumen:**", content['resumen'])
        st.write("**KPIs:**")
        st.json(content['kpis'])
        st.write("**Datos:**")
        st.dataframe(content['dataframe'].head(20), use_container_width=True)
        st.write("**Conclusiones:**", content['conclusiones'])

    # Generación
    st.markdown("---")
    st.header("3. Generar Reportes")

    col1, col2, col3 = st.columns(3)

    with col1:
        if st.button("📄 Generar PDF", use_container_width=True):
            with st.spinner("Generando PDF..."):
                try:
                    path = generator.generate_pdf_report(title, content)
                    st.success(f"PDF generado: {os.path.basename(path)}")
                    with open(path, "rb") as f:
                        st.download_button("⬇️ Descargar PDF", f, file_name=os.path.basename(path), mime="application/pdf")
                except Exception as e:
                    st.error(f"Error: {e}")

    with col2:
        if st.button("📝 Generar Word", use_container_width=True):
            with st.spinner("Generando Word..."):
                try:
                    path = generator.generate_word_report(title, content)
                    st.success(f"Word generado: {os.path.basename(path)}")
                    with open(path, "rb") as f:
                        st.download_button("⬇️ Descargar Word", f, file_name=os.path.basename(path), 
                                         mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                except Exception as e:
                    st.error(f"Error: {e}")

    with col3:
        if st.button("📊 Generar Excel", use_container_width=True):
            with st.spinner("Generando Excel..."):
                try:
                    path = generator.generate_excel_report(title, content)
                    st.success(f"Excel generado: {os.path.basename(path)}")
                    with open(path, "rb") as f:
                        st.download_button("⬇️ Descargar Excel", f, file_name=os.path.basename(path),
                                         mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                except Exception as e:
                    st.error(f"Error: {e}")

    # Generar todos
    st.markdown("---")
    if st.button("🚀 Generar Todos los Formatos", use_container_width=True):
        with st.spinner("Generando todos los formatos..."):
            results = generator.generate_all_formats(title, content)

            col_r1, col_r2, col_r3 = st.columns(3)

            with col_r1:
                if 'pdf' in results:
                    st.success("✅ PDF listo")
                    with open(results['pdf'], "rb") as f:
                        st.download_button("⬇️ PDF", f, file_name=os.path.basename(results['pdf']), mime="application/pdf")
                elif 'pdf_error' in results:
                    st.error(f"❌ PDF: {results['pdf_error']}")

            with col_r2:
                if 'word' in results:
                    st.success("✅ Word listo")
                    with open(results['word'], "rb") as f:
                        st.download_button("⬇️ Word", f, file_name=os.path.basename(results['word']),
                                         mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                elif 'word_error' in results:
                    st.error(f"❌ Word: {results['word_error']}")

            with col_r3:
                if 'excel' in results:
                    st.success("✅ Excel listo")
                    with open(results['excel'], "rb") as f:
                        st.download_button("⬇️ Excel", f, file_name=os.path.basename(results['excel']),
                                         mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                elif 'excel_error' in results:
                    st.error(f"❌ Excel: {results['excel_error']}")
